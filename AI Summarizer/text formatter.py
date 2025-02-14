import spacy
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Load SpaCy model
nlp = spacy.load("en_core_web_sm")

# Function to read text from .docx file
def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)

# Function to save structured text to a .docx file
def save_docx(text, output_path):
    doc = Document()
    for line in text.split("\n"):
        para = doc.add_paragraph(line)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        if para.runs:  # Check if there are any runs in the paragraph
            run = para.runs[0]
            run.font.size = Pt(12)
    doc.save(output_path)

# Path to your .docx file
file_path = "Summaries/2406 summary.docx"
output_path = "structured_output.docx"

# Read text from the .docx file
text = read_docx(file_path)

# Process the text with SpaCy
doc = nlp(text)

# Example: Structuring text with Named Entities highlighted
structured_text = ""
for sent in doc.sents:
    structured_text += f"{sent.text}\n"
    for ent in sent.ents:
        structured_text += f" - {ent.text} ({ent.label_})\n"

# Save structured text to a new .docx file
save_docx(structured_text, output_path)
